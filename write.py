import docx
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from take_output import details

HEADING = 24
QUES = 14
CODE = 12
details_printf = f'printf("{details}");'


def writeBase(doc, text, font_size, font_family = 'Roborto'):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font = run.font
    font.name = font_family
    font.size = Pt(font_size)
    return para, run

def writeHeading(doc, text):
    para, run = writeBase(doc, text, HEADING)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def writeQues(doc, text):
    para, run = writeBase(doc, text, QUES)
    para.paragraph_format.left_indent = Inches(-0.55)
    para.paragraph_format.right_indent = Inches(-0.55)

def writeCode(doc, file_name):
    text = readCodeFile(file_name)
    para, run = writeBase(doc, text, CODE)

def readCodeFile(file_name):
    with open(file_name) as f:
        text = f.read()
        f.close()

    text = text.replace('//add details', details_printf)
    return text

def addOutput(doc, image_name):
    para, run = writeBase(doc, None, 0)
    para.paragraph_format.left_indent = Inches(-1)
    run.add_picture(image_name, height = Inches(48/12), width = Inches(96/12))


